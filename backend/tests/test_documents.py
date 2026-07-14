from pathlib import Path

import fitz
from docx import Document

from app.documents import parse_document


def test_markdown_preserves_heading_and_paragraph(tmp_path: Path):
    path = tmp_path / "note.md"
    path.write_text("# 所得税\n\n递延所得税资产应满足确认条件。", encoding="utf-8")
    chunks = parse_document(path, ".md")
    assert chunks[0]["heading"] == "所得税"
    assert chunks[1]["locator"] == "第 2 段"


def test_docx_preserves_paragraph_locator(tmp_path: Path):
    path = tmp_path / "handout.docx"
    doc = Document()
    doc.add_heading("收入", level=1)
    doc.add_paragraph("收入确认需要满足相应条件。")
    doc.save(path)
    chunks = parse_document(path, ".docx")
    assert chunks[-1]["locator"] == "第 2 段"
    assert chunks[-1]["heading"] == "收入"


def test_pdf_preserves_page_number(tmp_path: Path):
    path = tmp_path / "book.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Audit evidence")
    pdf.save(path)
    chunks = parse_document(path, ".pdf")
    assert chunks[0]["locator"] == "第 1 页"
