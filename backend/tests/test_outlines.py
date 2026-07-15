from pathlib import Path

import fitz
from docx import Document as DocxDocument
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import sessionmaker

import app.outlines as outlines
from app.database import Base
from app.main import confirm_outline_proposal
from app.models import (
    Chapter, CurriculumProposal, CurriculumSourceLink, Document, DocumentChunk,
    ExamTrack, KnowledgePoint, ProposalNode, Subject,
)


def test_extracts_markdown_and_docx_heading_levels(tmp_path: Path):
    markdown = tmp_path / "tax.md"
    markdown.write_text("# 增值税\n\n## 征税范围\n\n正文", encoding="utf-8")
    headings = outlines.extract_headings(markdown, ".md")
    assert [(item.level, item.title) for item in headings] == [(1, "增值税"), (2, "征税范围")]

    docx_path = tmp_path / "accounting.docx"
    docx = DocxDocument()
    docx.add_heading("收入", level=1)
    docx.add_heading("收入确认", level=2)
    docx.save(docx_path)
    headings = outlines.extract_headings(docx_path, ".docx")
    assert [item.title for item in headings] == ["收入", "收入确认"]


def test_pdf_bookmarks_are_preferred(tmp_path: Path):
    path = tmp_path / "audit.pdf"
    pdf = fitz.open()
    pdf.new_page(); pdf.new_page()
    pdf.set_toc([[1, "审计证据", 1], [2, "函证", 2]])
    pdf.save(path)
    headings = outlines.extract_headings(path, ".pdf")
    assert [(item.level, item.title, item.locator) for item in headings] == [
        (1, "审计证据", "第 1 页"), (2, "函证", "第 2 页")
    ]


def test_mechanical_docx_outline_links_heading_body_until_next_section(tmp_path: Path):
    path = tmp_path / "income.docx"
    docx = DocxDocument()
    docx.add_heading("收入", level=1)
    docx.add_paragraph("收入章节总述。")
    docx.add_heading("收入确认", level=2)
    docx.add_paragraph("确认条件正文。")
    docx.add_heading("金融资产", level=1)
    docx.save(path)
    document = Document(name="讲义", original_name=path.name, path=str(path), mime_type="application/docx")
    chunks = [
        DocumentChunk(id=1, document_id=1, content="收入", locator="第 1 段", position=0, embedding=[]),
        DocumentChunk(id=2, document_id=1, content="收入章节总述。", locator="第 2 段", position=1, embedding=[]),
        DocumentChunk(id=3, document_id=1, content="收入确认", locator="第 3 段", position=2, embedding=[]),
        DocumentChunk(id=4, document_id=1, content="确认条件正文。", locator="第 4 段", position=3, embedding=[]),
        DocumentChunk(id=5, document_id=1, content="金融资产", locator="第 5 段", position=4, embedding=[]),
    ]
    tree = outlines.mechanical_tree(document, chunks)
    assert tree[0]["source_chunk_ids"] == [1, 2, 3, 4]
    assert tree[0]["points"][0]["source_chunk_ids"] == [3, 4]


def test_proposal_merge_and_confirm_preserve_mastery(tmp_path: Path, monkeypatch):
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine, expire_on_commit=False)
    monkeypatch.setattr(outlines, "SessionLocal", Session)
    monkeypatch.setattr(outlines, "_enhance_with_deepseek", lambda db, document, tree, chunks: (tree, None, None))
    db = Session()
    exam = ExamTrack(name="注册会计师", code="CPA")
    db.add(exam); db.flush()
    subject = Subject(exam_id=exam.id, name="会计")
    db.add(subject); db.flush()
    chapter = Chapter(subject_id=subject.id, name="收入", position=0)
    db.add(chapter); db.flush()
    point = KnowledgePoint(chapter_id=chapter.id, name="收入确认", mastery=72, status="reviewing")
    db.add(point); db.flush()
    path = tmp_path / "income.md"
    path.write_text("# 收入\n\n## 收入确认\n\n收入确认需要满足资料列示条件。", encoding="utf-8")
    document = Document(name="收入讲义", original_name="income.md", path=str(path), mime_type="text/markdown",
                        status="ready", parse_status="ready", outline_status="extracting", subject_id=subject.id, exam_id=exam.id)
    db.add(document); db.flush()
    db.add_all([
        DocumentChunk(document_id=document.id, content="# 收入", locator="第 1 段", position=0, embedding=[]),
        DocumentChunk(document_id=document.id, content="## 收入确认\n收入确认需要满足资料列示条件。", locator="第 2 段", position=1, embedding=[]),
    ])
    proposal = CurriculumProposal(document_id=document.id, subject_id=subject.id, status="extracting")
    db.add(proposal); db.commit()

    outlines.process_outline(proposal.id)
    db.expire_all()
    nodes = db.scalars(select(ProposalNode).where(ProposalNode.proposal_id == proposal.id)).all()
    assert all(node.action == "merge" for node in nodes)

    result = confirm_outline_proposal(proposal.id, db)
    assert result["ok"] is True
    assert db.get(KnowledgePoint, point.id).mastery == 72
    assert db.scalar(select(func.count()).select_from(CurriculumSourceLink)) == 3
    assert confirm_outline_proposal(proposal.id, db)["idempotent"] is True
